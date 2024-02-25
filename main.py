from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def parse_inline_style(style_str):
    styles = {}
    for part in style_str.split(';'):
        if ':' in part:
            key, value = part.split(':', 1)
            styles[key.strip().lower()] = value.strip()
    return styles

def apply_styles(run, styles):
    if 'font-size' in styles:
        size = re.match(r'(\d+)px', styles['font-size'])
        if size:
            run.font.size = Pt(int(size.group(1)))

    if 'font-family' in styles:
        run.font.name = styles['font-family']

    if 'font-weight' in styles and styles['font-weight'] == 'bold':
        run.font.bold = True

    if 'color' in styles:
        color = re.match(r'#([0-9a-fA-F]{6})', styles['color'])
        if color:
            run.font.color.rgb = RGBColor.from_string(color.group(1))

def html_to_word(html_content, output_file):
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    for element in soup.body.find_all(recursive=True):
        if isinstance(element, NavigableString):
            continue

        styles = parse_inline_style(element.get('style', ''))

        # Handle paragraphs and headers
        if element.name in ['h1', 'h2', 'h3', 'p']:
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            apply_styles(run, styles)

        # Handle list items
        elif element.name == 'li':
            parent = element.find_parent(['ul', 'ol'])
            if parent.name == 'ul':
                p = doc.add_paragraph(style='List Bullet')
            elif parent.name == 'ol':
                p = doc.add_paragraph(style='List Number')
            run = p.add_run(element.get_text())
            apply_styles(run, styles)

    doc.save(output_file)

# Example HTML content
html_content = """
Insert HTML content in here
"""

# Convert the HTML to a Word document
html_to_word(html_content, 'output.docx')